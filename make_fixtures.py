"""
Fixture generator — creates all test PDFs and DOCX files needed by the test suite.
Run once before running pytest: python make_fixtures.py
"""

from pathlib import Path
import sys

FIXTURES = Path(__file__).parent / "fixtures"
FIXTURES.mkdir(exist_ok=True)
(FIXTURES / "pdf").mkdir(exist_ok=True)
(FIXTURES / "docx").mkdir(exist_ok=True)

# ── PDF generation via reportlab ─────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("WARNING: reportlab not installed — PDF fixtures will be skipped.")
    print("  Run: pip install reportlab --break-system-packages")


def _make_pdf(filename: str, body_text: str, bib_text: str):
    if not HAS_REPORTLAB:
        return
    path = FIXTURES / "pdf" / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 60
    y = height - margin
    line_height = 14

    def _write_block(text: str):
        nonlocal y
        for line in text.split("\n"):
            # Word-wrap long lines
            while len(line) > 90:
                c.drawString(margin, y, line[:90])
                line = line[90:]
                y -= line_height
                if y < margin:
                    c.showPage()
                    y = height - margin
            c.drawString(margin, y, line)
            y -= line_height
            if y < margin:
                c.showPage()
                y = height - margin

    c.setFont("Helvetica-Bold", 11)
    _write_block("LNI Test Paper — " + filename)
    c.setFont("Helvetica", 9)
    _write_block("\n" + body_text)
    c.setFont("Helvetica-Bold", 10)
    _write_block("\nLiteraturverzeichnis\n")
    c.setFont("Helvetica", 9)
    _write_block(bib_text)
    c.save()
    print(f"  Created {path}")


# ── DOCX generation via python-docx ──────────────────────────────────────────
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed — DOCX fixtures will be skipped.")


def _make_docx(filename: str, body_text: str, bib_text: str):
    if not HAS_DOCX:
        return
    path = FIXTURES / "docx" / filename
    doc = Document()
    doc.add_heading("LNI Test Paper — " + filename, level=1)
    for line in body_text.split("\n"):
        doc.add_paragraph(line)
    doc.add_heading("Literaturverzeichnis", level=2)
    for line in bib_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(str(path))
    print(f"  Created {path}")


# =============================================================================
# FIXTURE DEFINITIONS
# Each entry: (filename, body_text, bib_text)
# =============================================================================

PDFS = [

    # ── F01: Happy path — perfectly formatted paper ──────────────────────────
    ("f01_perfect.pdf",
     "This paper explores neural networks [LBH15]. We extend the work of "
     "Vaswani et al. [VSP17]. The dataset was introduced in [De12].",
     """[LBH15] LeCun, Yann; Bengio, Yoshua; Hinton, Geoffrey: Deep Learning. In: Nature, Vol. 521, 2015; S. 436--444.
[VSP17] Vaswani, Ashish; Shazeer, Noam; Parmar, Niki; Uszkoreit, Jakob; Jones, Llion; Gomez, Aidan N.; Kaiser, Lukasz; Polosukhin, Illia: Attention Is All You Need. In: Advances in Neural Information Processing Systems, 2017; S. 5998--6008.
[De12] Deng, Li: The MNIST Database of Handwritten Digit Images for Machine Learning Research. In: IEEE Signal Processing Magazine, Vol. 29, No. 6, 2012; S. 141--142."""),

    # ── F02: Orphaned entry (in bib, never cited in body) ────────────────────
    ("f02_orphaned_entry.pdf",
     "We use transformers [VSP17] as the base model.",
     """[VSP17] Vaswani, Ashish; Shazeer, Noam; Parmar, Niki: Attention Is All You Need. In: Advances in Neural Information Processing Systems, 2017; S. 5998--6008.
[LBH15] LeCun, Yann; Bengio, Yoshua; Hinton, Geoffrey: Deep Learning. In: Nature, Vol. 521, 2015; S. 436--444."""),

    # ── F03: Missing bib entry (cited in body, not in bib) ───────────────────
    ("f03_missing_entry.pdf",
     "We use transformers [VSP17] and also reference BERT [Dev18].",
     """[VSP17] Vaswani, Ashish; Shazeer, Noam: Attention Is All You Need. In: Advances in Neural Information Processing Systems, 2017; S. 5998--6008."""),

    # ── F04: Wrong LNI key format ─────────────────────────────────────────────
    ("f04_bad_key_format.pdf",
     "The model [vaswani17] performs well.",
     """[vaswani17] Vaswani, Ashish: Attention Is All You Need. In: NeurIPS, 2017; S. 1--10.
[X] Smith, John: A Book. Springer, 2020."""),

    # ── F05: Single-dash page range (should be double dash) ──────────────────
    ("f05_single_dash_pages.pdf",
     "Results are in [AB20].",
     """[AB20] Autor, Bob; Becker, Alice: A Fine Paper. In: Journal of Computing, Vol. 12, 2020; S. 10-20."""),

    # ── F06: Wrong author order (Firstname Lastname instead of Lastname, Firstname) ──
    ("f06_author_order.pdf",
     "See [Jo21] for details.",
     """[Jo21] John Smith: Machine Learning Basics. In: IEEE Transactions, Vol. 5, 2021; S. 1--10."""),

    # ── F07: Future year ──────────────────────────────────────────────────────
    ("f07_future_year.pdf",
     "This will be published [Fu99].",
     """[Fu99] Future, Author: Paper From The Future. Springer, 2099."""),

    # ── F08: Key inconsistency (initials mismatch) ────────────────────────────
    ("f08_key_mismatch.pdf",
     "Results from [XY15] confirm the hypothesis.",
     """[XY15] Mueller, Hans; Schmidt, Klaus: Algorithmic Methods. Springer, 2020."""),

    # ── F09: Missing required fields (article missing journal+pages) ──────────
    ("f09_missing_fields.pdf",
     "The algorithm [No20] runs fast.",
     """[No20] Nobody, Anon: A Paper Without Fields. 2020."""),

    # ── F10: Real paper with DOI (for verification) ────────────────────────────
    ("f10_doi_verification.pdf",
     "Deep learning has transformed NLP [Va17].",
     """[Va17] Vaswani, Ashish; Shazeer, Noam; Parmar, Niki: Attention Is All You Need. In: Advances in Neural Information Processing Systems, 2017; S. 5998--6008. doi: 10.48550/arXiv.1706.03762"""),

    # ── F11: Hallucinated / fake reference ────────────────────────────────────
    ("f11_fake_reference.pdf",
     "As shown in [Fa23], our approach outperforms all baselines.",
     """[Fa23] FakeAuthor, John; NonExistent, Jane: Revolutionary Quantum AI Blockchain Method That Does Not Exist At All. In: Journal of Imaginary Research, Vol. 999, 2023; S. 1--500."""),

    # ── F12: German bibliography heading ──────────────────────────────────────
    ("f12_german_heading.pdf",
     "Wie in [Mu20] gezeigt wird.",
     "Quellenverzeichnis\n[Mu20] Mueller, Hans: Informatik Grundlagen. Springer, 2020."),

    # ── F13: Numeric citation keys [1], [2], [3] ─────────────────────────────
    ("f13_numeric_keys.pdf",
     "Deep learning [1] and transformers [2] are powerful methods.",
     """[1] LeCun, Yann; Bengio, Yoshua: Deep Learning. In: Nature, Vol. 521, 2015; S. 436--444.
[2] Vaswani, Ashish: Attention Is All You Need. In: NeurIPS, 2017; S. 5998--6008."""),

    # ── F14: Website citation with urldate ───────────────────────────────────
    ("f14_website_citation.pdf",
     "The official LNI style guide [GI24] is available online.",
     """[GI24] Gesellschaft fuer Informatik: LNI Style Guide. https://gi.de/lni Stand: 15.05.2024."""),

    # ── F15: Website citation WITHOUT urldate (should flag missing field) ────
    ("f15_website_no_urldate.pdf",
     "See [We22] for the documentation.",
     """[We22] Some Author: Documentation Page. https://example.com/docs"""),

    # ── F16: Same author cited multiple times, year disambiguation [Wa14a] ───
    ("f16_disambiguation.pdf",
     "First work [Wa14a] and second work [Wa14b] from same author.",
     """[Wa14a] Wagner, Klaus: First Paper. In: IEEE Transactions, Vol. 5, 2014; S. 1--10.
[Wa14b] Wagner, Klaus: Second Paper. In: IEEE Transactions, Vol. 6, 2014; S. 11--20."""),

    # ── F17: Very long author list (et al. key) ────────────────────────────────
    ("f17_long_author_list.pdf",
     "The BERT model [De18] achieved strong results.",
     """[De18] Devlin, Jacob; Chang, Ming-Wei; Lee, Kenton; Toutanova, Kristina: BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. In: Proceedings of NAACL-HLT, 2019; S. 4171--4186."""),

    # ── F18: BibTeX crossref resolution ──────────────────────────────────────
    ("f18_crossref_bib.pdf",
     "The method [AB20] builds on prior work.",
     """[AB20] Author, Bob: A Paper In A Proceedings. In: Proc. Workshop on Methods, 2020; S. 1--8."""),

    # ── F19: Paper body with no bibliography section at all ──────────────────
    ("f19_no_bibliography.pdf",
     "This paper has no references at all. It just makes claims without citations.",
     ""),

    # ── F20: Mixed correct and incorrect entries (batch test) ─────────────────
    ("f20_mixed_batch.pdf",
     "Using [LBH15] for DL, [FA99] for myth, and [VSP17] for attention.",
     """[LBH15] LeCun, Yann; Bengio, Yoshua; Hinton, Geoffrey: Deep Learning. In: Nature, Vol. 521, 2015; S. 436--444.
[FA99] FakeAuthor, XYZ: Nonexistent Paper About Nothing Real. In: Fake Journal, Vol. 0, 1999; S. 1-5.
[VSP17] Vaswani, Ashish; Shazeer, Noam: Attention Is All You Need. In: NeurIPS, 2017; S. 5998--6008."""),

    # ── F21: Implausible page span (>200 pages for single article) ────────────
    ("f21_huge_page_span.pdf",
     "The survey [Su10] covers everything.",
     """[Su10] Survey, Author: A Survey Article Claiming 500 Pages. In: Journal of Computing, Vol. 1, 2010; S. 1--501."""),

    # ── F22: F. Lastname author format (initial before surname) ──────────────
    ("f22_initial_author.pdf",
     "Results from [SM15] are significant.",
     """[SM15] J. Smith; K. Miller: Neural Approach to Problem Solving. In: IEEE Transactions, Vol. 10, 2015; S. 100--110."""),

    # ── F23: Empty bibliography section ──────────────────────────────────────
    ("f23_empty_bibliography.pdf",
     "The paper cites [AB20] in the body.",
     "Literaturverzeichnis"),

    # ── F24: Multiple citation styles mixed in one paper ─────────────────────
    ("f24_mixed_citation_styles.pdf",
     "See [VSP17] for attention and (LeCun 2015) for a different style.",
     """[VSP17] Vaswani, Ashish: Attention Is All You Need. In: NeurIPS, 2017; S. 1--10."""),

    # ── F25: Unicode/Umlaut authors ──────────────────────────────────────────
    ("f25_unicode_authors.pdf",
     "Research by [MÖ20] shows improvements.",
     """[MO20] Müller, Jörg; Özdemir, Ayşe: Verbesserungen im maschinellen Lernen. In: Informatik Spektrum, Vol. 43, 2020; S. 100--110."""),

]

DOCXS = [
    # ── D01: Perfect DOCX ────────────────────────────────────────────────────
    ("d01_perfect.docx",
     "This paper uses deep learning [LBH15] as the base.",
     """[LBH15] LeCun, Yann; Bengio, Yoshua; Hinton, Geoffrey: Deep Learning. In: Nature, Vol. 521, 2015; S. 436--444."""),

    # ── D02: DOCX with orphaned + missing entries ─────────────────────────────
    ("d02_cross_check.docx",
     "See [AA01] in the body but not the bib. [BB02] is orphaned.",
     """[BB02] Orphaned, Author: An Unused Reference. Springer, 2002.
[CC03] Cited, Author: Used Reference. Springer, 2003."""),
]


def main():
    print("Generating PDF fixtures...")
    for fname, body, bib in PDFS:
        _make_pdf(fname, body, bib)

    print("\nGenerating DOCX fixtures...")
    for fname, body, bib in DOCXS:
        _make_docx(fname, body, bib)

    # Also write raw text fixtures for parser unit tests (no PDF dependency)
    print("\nGenerating plain-text bib fixtures...")
    _write_txt_fixtures()

    print("\nAll fixtures created.")


def _write_txt_fixtures():
    """Write raw bib-section text fixtures for unit tests that don't need a PDF."""
    txt_dir = FIXTURES / "txt"
    txt_dir.mkdir(exist_ok=True)

    cases = {
        "bib_perfect.txt": (
            "[LBH15] LeCun, Yann; Bengio, Yoshua; Hinton, Geoffrey: Deep Learning. "
            "In: Nature, Vol. 521, 2015; S. 436--444.\n"
            "[VSP17] Vaswani, Ashish; Shazeer, Noam: Attention Is All You Need. "
            "In: NeurIPS, 2017; S. 5998--6008."
        ),
        "bib_bad_keys.txt": (
            "[vaswani2017] Vaswani, Ashish: Attention. In: NeurIPS, 2017; S. 1--10.\n"
            "[X] Smith, John: A Book. Springer, 2020.\n"
            "[TOOLONGKEY99] Nobody: Test. Springer, 1999."
        ),
        "bib_single_dash.txt": (
            "[AB20] Author, Bob: Good Paper. In: Journal, Vol. 1, 2020; S. 5-15."
        ),
        "bib_author_order.txt": (
            "[Jo21] John Smith: Bad Format. In: Journal, Vol. 1, 2021; S. 1--5."
        ),
        "bib_future_year.txt": (
            "[Fu99] Future, Author: Future Paper. Springer, 2099."
        ),
        "bib_key_mismatch.txt": (
            "[XY15] Mueller, Hans; Schmidt, Klaus: A Paper. Springer, 2020."
        ),
        "bib_website_no_urldate.txt": (
            "[We22] Author, John: A Website. https://example.com"
        ),
        "bib_website_with_urldate.txt": (
            "[GI24] GI: LNI Guide. https://gi.de/lni Stand: 15.05.2024."
        ),
        "bib_disambiguation.txt": (
            "[Wa14a] Wagner, Klaus: First Paper. In: IEEE, Vol. 5, 2014; S. 1--10.\n"
            "[Wa14b] Wagner, Klaus: Second Paper. In: IEEE, Vol. 6, 2014; S. 11--20."
        ),
        "bib_huge_page_span.txt": (
            "[Su10] Survey, Author: Big Survey. In: Journal, Vol. 1, 2010; S. 1--501."
        ),
        "bib_doi_entry.txt": (
            "[Va17] Vaswani, Ashish: Attention Is All You Need. In: NeurIPS, 2017. "
            "doi: 10.48550/arXiv.1706.03762"
        ),
        "bib_isbn_entry.txt": (
            "[Go09] Goodfellow, Ian: Deep Learning. MIT Press, 2016. ISBN 978-0-262-03561-3"
        ),
        "bib_numeric_keys.txt": (
            "[1] LeCun, Yann: Deep Learning. In: Nature, Vol. 521, 2015; S. 436--444.\n"
            "[2] Vaswani, Ashish: Attention. In: NeurIPS, 2017; S. 1--10.\n"
            "[10] Smith, John: A Book. Springer, 2020."
        ),
        "bib_german_umlaut.txt": (
            "[MU20] Müller, Jörg; Özdemir, Ayşe: Maschinelles Lernen. "
            "In: Informatik Spektrum, Vol. 43, 2020; S. 100--110."
        ),
        "bib_initial_author.txt": (
            "[SM15] J. Smith; K. Miller: Neural Approach. In: IEEE Trans., Vol. 10, 2015; S. 100--110."
        ),
        "bib_empty.txt": "",
        "bib_multiline.txt": (
            "[LBH15] LeCun, Yann; Bengio, Yoshua;\nHinton, Geoffrey: Deep Learning.\n"
            "In: Nature, Vol. 521,\n2015; S. 436--444."
        ),
    }

    for fname, content in cases.items():
        path = txt_dir / fname
        path.write_text(content, encoding="utf-8")
        print(f"  Created {path}")


if __name__ == "__main__":
    main()
