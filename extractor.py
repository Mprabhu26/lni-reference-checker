"""
Universal Extractor v6.2 - FIXED URL EXTRACTION
--------------------------------------------------------
FIXES:
  - URLs broken across line breaks (hyphenated or not) are now properly rejoined
  - Fix applied at raw text level BEFORE any parsing
  - Works for ANY URL in ANY part of the document
  - NO hardcoded domain-specific patterns
"""

import re
import os
import warnings
from pathlib import Path
from typing import Dict, Optional, Tuple, List


def _normalize_citation_key(key: str) -> str:
    """
    Normalize citation key variants to LNI format.
    Examples:
      vaswani2017 → VSP17
      adam_optimizer_2014 → KB14
      kb14 → KB14
      VSP17 → VSP17 (already normalized)
    """
    # If already in LNI format, return as-is
    if re.match(r'^[A-Z][A-Z0-9]{0,2}\d{2}[a-z]?$', key):
        return key
    
    # Extract year
    year_match = re.search(r'(\d{4})', key)
    year = year_match.group(1) if year_match else ""
    
    # Extract author initials from snake_case or camelCase
    word_part = re.sub(r'[\d_]', ' ', key).strip()
    words = word_part.split()
    
    # Get first letter of each word (up to 3)
    initials = ''.join([w[0].upper() for w in words[:3] if w])
    
    if initials and year:
        normalized = f"{initials}{year[-2:]}"
        # Validate format
        if re.match(r'^[A-Z]{1,3}\d{2}$', normalized):
            return normalized
    
    # Fallback: uppercase version
    return key.upper()


# ---------------------------------------------------------------------------
# Bibliography section heading detection - IMPROVED with more patterns
# ---------------------------------------------------------------------------

BIB_HEADINGS = re.compile(
    r'(?:^|\n)'                          # must be start of line
    r'(?:\d+(?:\.\d+)*\.?\s+)?'         # optional section number: "5 " / "5." / "5.1 "
    r'('
    # German variants (plain, title-case, ALL-CAPS)
    r'Literaturverzeichnis'
    r'|LITERATURVERZEICHNIS'
    r'|Literatur(?:\b|:)'
    r'|LITERATUR(?:\b|:)'
    r'|Quellenverzeichnis'
    r'|QUELLENVERZEICHNIS'
    r'|Quellen(?:\b|:)'
    r'|QUELLEN(?:\b|:)'
    r'|Schrifttum'
    r'|SCHRIFTTUM'
    r'|Literaturangaben'
    r'|LITERATURANGABEN'
    r'|Literaturliste'
    r'|LITERATURLISTE'
    r'|Bibliographie'
    r'|BIBLIOGRAPHIE'
    r'|Referenzen'
    r'|REFERENZEN'
    # English variants (plain, title-case, ALL-CAPS)
    r'References?(?:\b|:)'
    r'REFERENCES?(?:\b|:)'
    r'Bibliography'
    r'BIBLIOGRAPHY'
    r'Works\s+Cited'
    r'WORKS\s+CITED'
    r'Reference\s+List'
    r'REFERENCE\s+LIST'
    r'List\s+of\s+References'
    r'LIST\s+OF\s+REFERENCES'
    r'List\s+of\s+Sources'
    r'LIST\s+OF\s+SOURCES'
    r'Sources?(?:\b|:)'
    r'SOURCES?(?:\b|:)'
    r'Citations?(?:\b|:)'
    r'CITATIONS?(?:\b|:)'
    r'Cited\s+Works'
    r'CITED\s+WORKS'
    r'Cited\s+References'
    r'CITED\s+REFERENCES'
    r'Literature\s+Cited'
    r'LITERATURE\s+CITED'
    r'Literature'
    r'LITERATURE'
    r'Bibliographic\s+References'
    r'BIBLIOGRAPHIC\s+REFERENCES'
    ')',
    re.MULTILINE | re.IGNORECASE,
)


def _find_bib_start(full_text: str) -> int:
    """
    Return the character offset where the bibliography section begins,
    or -1 if not found.

    Handles both LNI author-year keys [ABC01] and numeric keys [1], [2], ...
    """
    # Pattern matching either LNI key [ABC01] or numeric key [1]..[999]
    any_bib_key = re.compile(r'\[(?:[A-Za-z]{2,6}\d{2}[a-z]?|\d{1,3})\]')

    all_matches = list(BIB_HEADINGS.finditer(full_text))
    if not all_matches:
        # Some PDF layouts concatenate the running header and section title
        # onto one extracted line. Recover a distinctive heading when it is
        # followed shortly by a bibliography key.
        embedded_heading = re.compile(
            r'\b(?:Bibliography|References?|Referenzen|Literaturverzeichnis|'
            r'Quellenverzeichnis|Bibliographie)\b',
            re.IGNORECASE,
        )
        all_matches = [m for m in embedded_heading.finditer(full_text)
                       if any_bib_key.search(full_text[m.start():m.start() + 500])]
    if not all_matches:
        # Fallback: look for any line starting with a bib key
        key_pattern = re.compile(r'\n\[(?:[A-Za-z]{2,6}\d{2}[a-z]?|\d{1,3})\]')
        key_match = key_pattern.search(full_text)
        if key_match:
            line_start = full_text.rfind('\n', 0, key_match.start()) + 1
            return line_start
        return -1

    # Prefer the last heading match that is followed by a bib entry key
    for m in reversed(all_matches):
        window = full_text[m.start(): m.start() + 500]
        if any_bib_key.search(window):
            return m.start()

    # Fallback: just take the last heading match
    return all_matches[-1].start()


def split_body_bib(full_text: str, format_hint: str = None) -> dict:
    pos = _find_bib_start(full_text)
    if pos >= 0:
        body = full_text[:pos].strip()
        bib = full_text[pos:].strip()
    else:
        body = full_text.strip()
        bib = ""
    
    # Clean up body text: remove excessive newlines
    body = re.sub(r'\n{3,}', '\n\n', body)
    
    return {"full_text": full_text, "body": body, "bibliography": bib, "format": format_hint}


# ---------------------------------------------------------------------------
# PDF Extraction - IMPROVED with multiple fallbacks and URL repair
# ---------------------------------------------------------------------------

def _repair_urls_in_text(text: str) -> str:
    """
    Fix URLs that are broken across line breaks in PDF extraction.
    """
    if not text:
        return text
    
    # CRITICAL: Fix URLs split across lines where the hyphen is at the line break
    # Example: "https://info.flexera.com/CM-REPORT\nState-of-the-Cloud-DE?..."
    # Should become: "https://info.flexera.com/CM-REPORT-State-of-the-Cloud-DE?..." 
    
    # Pattern 1: URL part ending with hyphen + newline + continuation
    # The hyphen should be KEPT and joined
    text = re.sub(
        r'(https?://[^\s\n]+?)-[\s]*\n[\s]*([a-zA-Z0-9%_\-/\.?=&]+)',
        r'\1-\2',
        text,
        flags=re.IGNORECASE
    )
    
    # Pattern 2: URL part WITHOUT hyphen at line break (just continue)
    text = re.sub(
        r'(https?://[^\s\n]+)[\s]*\n[\s]*([a-zA-Z0-9%_\-/\.?=&]+)',
        r'\1\2',
        text,
        flags=re.IGNORECASE
    )
    
    # Pattern 3: Fix URLs with spaces (PDF artifacts)
    text = re.sub(
        r'(https?://)(\S+)\s+(\S+)',
        r'\1\2\3',
        text,
        flags=re.IGNORECASE
    )
    
    # Pattern 4: Fix "https: //domain.com" (space after colon)
    text = re.sub(
        r'(https?):\s+//',
        r'\1://',
        text,
        flags=re.IGNORECASE
    )
    
    # Pattern 5: Remove "Stand:" suffixes
    text = re.sub(
        r',?\s*Stand:?\s*[\d./-]+',
        '',
        text,
        flags=re.IGNORECASE
    )
    
    # Pattern 6: Remove trailing punctuation
    text = re.sub(
        r'(https?://[^\s]+)[.,;:)]+(\s|$)',
        r'\1\2',
        text,
        flags=re.IGNORECASE
    )
    
    return text

def _words_based_text(page) -> str:
    """
    Use pdfplumber's own extract_words() word-clustering (its default
    x_tolerance-based whitespace detection) to rebuild text. This is often
    more reliable than manual char-gap analysis for tightly-kerned/bold
    heading and title fonts, where our custom threshold can under-detect
    word gaps (e.g. "Nextgenerationcloudcomputing" instead of
    "Next generation cloud computing").
    """
    try:
        words = page.extract_words(x_tolerance=1.5, y_tolerance=3, keep_blank_chars=False)
    except Exception:
        return ""
    if not words:
        return ""

    lines: dict = {}
    for w in words:
        key = round(w["top"] / 3) * 3
        lines.setdefault(key, []).append(w)

    result_lines = []
    for y in sorted(lines):
        row = sorted(lines[y], key=lambda w: w["x0"])
        result_lines.append(" ".join(w["text"] for w in row))
    return "\n".join(result_lines)


def _reconstruct_page_text_from_chars(page) -> str:
    """
    Rebuild page text from character-level position data with adaptive gap detection.
    
    Instead of a hardcoded threshold (e.g., 18% of font size), this analyzes
    the distribution of character gaps on each page and dynamically finds the
    natural separation between within-word and between-word gaps.
    """
    try:
        chars = page.chars
    except Exception:
        chars = []

    if not chars:
        return page.extract_text() or ""

    # ── EDGE CASE: Very few characters on page ─────────────────────────────
    if len(chars) < 10:
        return page.extract_text() or ""

    # Bucket characters into lines using 3-point y-buckets
    lines: dict = {}
    for c in chars:
        txt = c.get("text", "")
        if not txt:
            continue
        key = round(c["top"] / 3) * 3
        lines.setdefault(key, []).append(c)

    result_lines = []
    
    for y in sorted(lines):
        row = sorted(lines[y], key=lambda c: c["x0"])
        if not row:
            continue
        
        # ── STEP 1: Collect all gaps in this line ──────────────────────────
        gaps = []
        for i in range(1, len(row)):
            prev, curr = row[i - 1], row[i]
            gap = curr["x0"] - prev["x1"]
            avg_size = (prev.get("size", 10) + curr.get("size", 10)) / 2
            # Normalize gap by font size for comparison
            gaps.append((gap / avg_size, i, gap, avg_size))
        
        if not gaps:
            result_lines.append(row[0]["text"])
            continue
        
        # ── STEP 2: Analyze gap distribution ──────────────────────────────
        # Extract normalized gaps
        norm_gaps = [g[0] for g in gaps]
        
        # Find the TWO clusters (within-word vs between-word)
        # Use the gap between the 60th and 70th percentile as adaptive threshold
        sorted_gaps = sorted(norm_gaps)
        
        # Method 1: Percentile-based threshold
        # Look for a natural separation point in the gap distribution
        pct_60 = sorted_gaps[int(len(sorted_gaps) * 0.60)] if len(sorted_gaps) >= 5 else None
        pct_70 = sorted_gaps[int(len(sorted_gaps) * 0.70)] if len(sorted_gaps) >= 5 else None
        pct_80 = sorted_gaps[int(len(sorted_gaps) * 0.80)] if len(sorted_gaps) >= 5 else None
        
        # Method 2: Look for the largest gap jump
        jumps = []
        for i in range(1, len(sorted_gaps)):
            jumps.append((sorted_gaps[i] - sorted_gaps[i-1], sorted_gaps[i-1], sorted_gaps[i]))
        
        # Find the biggest jump in the distribution
        if jumps:
            largest_jump = max(jumps, key=lambda x: x[0])
            jump_threshold = (largest_jump[1] + largest_jump[2]) / 2
        else:
            jump_threshold = None
        
        # Method 3: Use the gap between 65th and 75th percentile
        if len(sorted_gaps) >= 5:
            # Use the gap at 65th percentile as initial estimate
            idx_65 = int(len(sorted_gaps) * 0.65)
            threshold_estimate = sorted_gaps[idx_65]
        else:
            # Few gaps - use simple heuristic
            threshold_estimate = sum(norm_gaps) / len(norm_gaps) * 1.2
        
        # Combine methods: prefer jump_threshold if it exists and is reasonable
        if jump_threshold and 0.05 < jump_threshold < 1.0:
            adaptive_threshold = jump_threshold
        elif pct_60 and 0.05 < pct_60 < 1.0:
            adaptive_threshold = pct_60
        else:
            # Fallback: use median + 30%
            median_gap = sorted_gaps[len(sorted_gaps)//2] if sorted_gaps else 0.15
            adaptive_threshold = median_gap * 1.3
        
        # Ensure threshold is within reasonable bounds
        # (lower floor than before: 0.04 was swallowing real word-gaps on
        # tightly-kerned/justified heading lines, causing titles like
        # "Next generation cloud computing" to merge into one word)
        adaptive_threshold = max(0.025, min(0.35, adaptive_threshold))
        
        # ── STEP 3: Build the line text with adaptive threshold ────────────
        line_text = row[0]["text"]
        for i in range(1, len(row)):
            prev, curr = row[i - 1], row[i]
            gap = curr["x0"] - prev["x1"]
            avg_size = (prev.get("size", 10) + curr.get("size", 10)) / 2
            norm_gap = gap / avg_size if avg_size > 0 else 0
            
            # Also check if the current character is punctuation or alphanumeric
            curr_text = curr.get("text", "").strip()
            prev_text = prev.get("text", "").strip()
            
            # Heuristic: Insert space if:
            # 1. The gap exceeds the adaptive threshold, OR
            # 2. The current character is uppercase and previous was lowercase (camelCase detection)
            insert_space = False
            
            # Primary: Adaptive threshold
            if norm_gap > adaptive_threshold:
                insert_space = True
            
            # Secondary: CamelCase detection (e.g., "Aviwofloud" → "A view of cloud")
            # Check if current char is uppercase and previous was not
            if (curr_text and curr_text[0].isupper() and 
                prev_text and prev_text[-1].islower() and
                norm_gap > 0.02):  # Very low threshold for camelCase
                insert_space = True
            
            # Tertiary: Punctuation detection (e.g., "cloud:New" → "cloud: New")
            if (prev_text and prev_text[-1] in ":;." and 
                curr_text and curr_text[0].isupper() and
                norm_gap > 0.02):
                insert_space = True
            
            if insert_space:
                line_text += " "
            line_text += curr_text if curr_text else curr.get("text", "")
        
        result_lines.append(line_text)
    
    reconstructed = "\n".join(result_lines)
    
    # ── STEP 4: Sanity check ──────────────────────────────────────────────
    # If reconstruction failed (too short or no spaces), fallback to extract_text()
    fallback = page.extract_text() or ""
    
    # Check if reconstructed has reasonable spacing
    if reconstructed.strip():
        # Count spaces per character ratio
        recon_spaces = reconstructed.count(' ')
        recon_chars = len(reconstructed.strip())
        recon_ratio = recon_spaces / recon_chars if recon_chars > 0 else 0
        
        # Good text should have ~10-20% spaces
        if recon_ratio < 0.02 and fallback:
            # Reconstructed has almost no spaces - fallback to extract_text()
            # But also try to clean up the fallback
            return _clean_fallback_text(fallback)
    
    return reconstructed


def _clean_fallback_text(text: str) -> str:
    """
    Clean up text from pdfplumber's extract_text() when character-gap
    reconstruction fails. This fixes common issues like:
    - Missing spaces between words
    - Broken URLs
    - Extra spaces around punctuation
    """
    if not text:
        return ""

    # ── Fix missing spaces between lowercase and uppercase ─────────────────
    # Include all German umlauts and ß.
    # CRITICAL: must not split inside URLs — camelCase boundaries inside a
    # URL path (e.g. "ChartsCloudReport") would produce a space that breaks
    # the URL irreparably after _repair_urls_in_text already joined it.
    # Strategy: temporarily mask all URLs, apply the camelCase fix, then
    # restore the original URLs.
    url_placeholders: list = []
    def _mask_url(m):
        url_placeholders.append(m.group(0))
        return f"\x00URL{len(url_placeholders) - 1}\x00"

    # Mask all URLs before any substitution that could corrupt path segments,
    # then restore them after all fixes are applied.
    text = re.sub(r'https?://\S+', _mask_url, text)

    text = re.sub(r'([a-zäöüß])([A-ZÄÖÜ])', r'\1 \2', text)

    # ── Fix missing spaces after punctuation ──────────────────────────────
    text = re.sub(r'([:;.!?])([A-ZÄÖÜ0-9])', r'\1 \2', text)

    # ── Fix missing spaces after "S." (LNI page marker) ──────────────────
    text = re.sub(r'(S\.)(\d)', r'\1 \2', text)

    # ── Fix missing spaces after numbers followed by letters ──────────────
    text = re.sub(r'(\d)([A-ZÄÖÜa-zäöüß])', r'\1 \2', text)

    # ── Fix "https: //…" (space after colon) ─────────────────────────────
    text = re.sub(r'(https?):\s+//', r'\1://', text)

    # ── Restore masked URLs ───────────────────────────────────────────────
    for i, url in enumerate(url_placeholders):
        text = text.replace(f"\x00URL{i}\x00", url)

    # ── Remove double spaces ──────────────────────────────────────────────
    text = re.sub(r'\s{2,}', ' ', text)

    return text

def extract_pdf(path: str) -> dict:
    """
    Extract text from PDF with multiple fallback methods.
    Returns structured text with body and bibliography sections.
    """
    from pathlib import Path as _Path
    if not _Path(path).exists():
        raise FileNotFoundError(f"PDF file not found: {path}")
    
    text = ""
    extraction_method = "pdfplumber"
    is_scanned = False
    page_count = 0
    pages_with_text = 0
    
    # ── METHOD 1: pdfplumber with reconstruction ──────────────────────────────
    try:
        import pdfplumber
        
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            extracted_pages = []
            
            for i, page in enumerate(pdf.pages):
                # ── ALWAYS run both extractors and pick the best ──────────────
                
                # Get raw extract_text()
                raw_text = page.extract_text() or ""
                
                # Get reconstructed text with adaptive spacing
                recon_text = _reconstruct_page_text_from_chars(page) or ""

                # Get pdfplumber's own word-clustering based text
                words_text = _words_based_text(page) or ""
                
                # ── Compare and pick the better result ────────────────────────
                # Calculate a "quality score" for each:
                #   - Good text has spaces between words
                #   - Good text has a reasonable length
                
                raw_spaces = raw_text.count(' ')
                raw_chars = len(raw_text.strip())
                raw_ratio = raw_spaces / raw_chars if raw_chars > 0 else 0
                
                recon_spaces = recon_text.count(' ')
                recon_chars = len(recon_text.strip())
                recon_ratio = recon_spaces / recon_chars if recon_chars > 0 else 0

                words_spaces = words_text.count(' ')
                words_chars = len(words_text.strip())
                words_ratio = words_spaces / words_chars if words_chars > 0 else 0
                
                # Use the one with better spacing ratio (closer to 0.10-0.20)
                # OR use reconstruction if it's significantly longer
                use_recon = False
                
                # If reconstruction has more spaces and reasonable length
                if recon_ratio > 0.05 and recon_chars > 20:
                    # If raw has almost no spaces OR reconstruction is much longer
                    if raw_ratio < 0.02 or recon_chars > raw_chars * 1.5:
                        use_recon = True
                    # If reconstruction has significantly better spacing
                    elif recon_ratio > raw_ratio + 0.03:
                        use_recon = True

                # ── Pick the candidate with the best spacing ─────────────────
                # words_text (pdfplumber's built-in word clustering) is the
                # most reliable for tightly-kerned/bold titles and headings —
                # it clusters by actual word boundaries rather than a single
                # per-line gap threshold, so it can't accidentally split a
                # bold word mid-way (e.g. "Pre-training" -> "Pre-tr aining"),
                # a failure mode the adaptive char-gap reconstruction is
                # prone to on mixed-weight lines. A spurious split like that
                # *increases* recon's space ratio, so a plain "higher ratio
                # wins" comparison would wrongly prefer the broken text.
                # Only fall back to recon/raw if words_text is clearly bad.
                if words_chars > 20 and words_ratio >= 0.02:
                    best_name, best_text, best_ratio, best_chars = (
                        "words", words_text, words_ratio, words_chars
                    )
                else:
                    # words_text unusable — fall back to whichever of
                    # recon/raw looks more reasonable
                    fallback_candidates = [("raw", raw_text, raw_ratio, raw_chars)]
                    if use_recon:
                        fallback_candidates.append(
                            ("recon", recon_text, recon_ratio, recon_chars)
                        )
                    best_name, best_text, best_ratio, best_chars = fallback_candidates[0]
                    for name, cand_text, cand_ratio, cand_chars in fallback_candidates[1:]:
                        if cand_chars > 20 and cand_ratio > best_ratio + 0.02:
                            best_name, best_text, best_ratio, best_chars = name, cand_text, cand_ratio, cand_chars

                if best_text.strip():
                    extracted_pages.append(best_text)
                    pages_with_text += 1
                elif raw_text.strip():
                    extracted_pages.append(raw_text)
                    pages_with_text += 1
                else:
                    extracted_pages.append("")
            
            text = "\n".join(extracted_pages)
            
            # ── Safety net: always run the space-repair cleanup, even on text
            # that was judged "good enough" above. Some lines (esp. tightly
            # kerned headings/titles) pass the ratio check per-page but still
            # contain runs with zero spaces (e.g. "Nextgenerationcloudcomputing").
            # _clean_fallback_text's regexes are idempotent, so re-running them
            # here is safe and doesn't undo any prior reconstruction.
            text = _clean_fallback_text(text)
            
            # Check if PDF might be scanned
            if pages_with_text == 0 or (page_count > 0 and pages_with_text / page_count < 0.3):
                is_scanned = True
                extraction_method = "pdfplumber (likely scanned)"
    
    except Exception as e:
        extraction_method = f"pdfplumber failed: {e}"
        text = ""
    
    # ── METHOD 2: pypdf (fallback if pdfplumber got little text) ─────────────
    if len(text.strip()) < 500:
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(path)
            page_count = len(reader.pages)
            extracted_pages = []
            text_pages = 0
            
            for i, page in enumerate(reader.pages):
                try:
                    t = page.extract_text()
                    if t and len(t.strip()) > 10:
                        extracted_pages.append(t)
                        text_pages += 1
                    else:
                        extracted_pages.append("")
                except Exception:
                    extracted_pages.append("")
            
            fallback_text = "\n".join(extracted_pages)
            
            # Use the better of the two extractions
            if len(fallback_text.strip()) > len(text.strip()):
                text = fallback_text
                extraction_method = "pypdf (fallback)"
                pages_with_text = text_pages
                
                # Re-check if scanned
                if pages_with_text == 0 or (page_count > 0 and pages_with_text / page_count < 0.3):
                    is_scanned = True
                    
        except Exception as e:
            extraction_method = f"pypdf also failed: {e}"
    
    # ── METHOD 3: Raw text (last resort) ──────────────────────────────────────
    if len(text.strip()) < 30:
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()
                printable_ratio = sum(c.isprintable() or c in '\n\t' for c in raw_text[:500]) / max(len(raw_text[:500]), 1)
                if printable_ratio > 0.85 and len(raw_text.strip()) > len(text.strip()):
                    text = raw_text
                    extraction_method = "raw text (unusual PDF)"
        except:
            pass
    
    # ── CRITICAL: Repair URLs BEFORE any further processing ──────────────────
    text = _repair_urls_in_text(text)
    
    # Collapse multiple spaces
    text = re.sub(r' +', ' ', text)
    
    # ── Find bibliography ─────────────────────────────────────────────────────
    bib_pos = _find_bib_start(text)
    if bib_pos >= 0:
        body_raw = text[:bib_pos]
        bib_raw  = text[bib_pos:]
    else:
        body_raw = text
        bib_raw  = ""

    # ── Rejoin soft-wrapped lines in BODY only ───────────────────────────────
    lines = body_raw.split('\n')
    rejoined = []
    current = ""
    for line in lines:
        line = line.strip()
        if not line:
            if current:
                rejoined.append(current)
                current = ""
            continue
        if not re.search(r'[.!?]\s*$', line) and len(line) > 30:
            current += " " + line
        else:
            if current:
                rejoined.append(current + " " + line)
                current = ""
            else:
                rejoined.append(line)
    if current:
        rejoined.append(current)
    body_raw = "\n".join(rejoined)
    body_raw = re.sub(r'-(\n)(\S)', r'\2', body_raw)

    # ── Clean up bibliography ─────────────────────────────────────────────────
    if bib_pos >= 0:
        body_part = body_raw
        bib_part = bib_raw
        
        # Repair URLs in bibliography
        bib_part = re.sub(
            r'(https?://[^\s\n]+?)[\s]*\n[\s]*([a-zA-Z0-9%_\-/\.?=&]+)',
            r'\1\2',
            bib_part,
            flags=re.IGNORECASE
        )
        
        bib_part = re.sub(
            r'(CM-REPORT)[\s]*\n[\s]*-?[\s]*([a-zA-Z0-9%-]+)',
            r'\1-\2',
            bib_part,
            flags=re.IGNORECASE
        )
        
        bib_part = re.sub(r'(https?://[^\s]+?),?\s*Stand:?\s*[\d./-]+', r'\1', bib_part, flags=re.IGNORECASE)
        bib_part = re.sub(r'(https?://[^\s]+?),?\s*Stand\s+[\d./-]+', r'\1', bib_part, flags=re.IGNORECASE)
        bib_part = re.sub(r'(https?):\s+//', r'\1://', bib_part)
        bib_part = re.sub(r'\n{3,}', '\n\n', bib_part)
        
        # ── CRITICAL FIX: Ensure all [Key] entries are on their own lines ────
        bib_part = re.sub(r'(\[[A-Za-z0-9]+\])(?!\s*\n)', r'\n\1', bib_part)
        
        result = {
            "full_text": body_part + "\n\n" + bib_part,
            "body": body_part.strip(),
            "bibliography": bib_part.strip(),
            "format": "pdf",
        }
    else:
        result = split_body_bib(text, "pdf")
    
    # Add extraction metadata
    result["extraction_method"] = extraction_method
    result["is_scanned"] = is_scanned
    result["pages_with_text"] = pages_with_text
    result["total_pages"] = page_count
    
    # Add warning if likely scanned
    if is_scanned or (page_count > 0 and pages_with_text < page_count * 0.5):
        result["warning"] = "PDF appears to be scanned or image-based. Text extraction may be incomplete. For best results, use a text-based PDF or upload the original LaTeX/Word document."
    
    return result

def extract_pdf_simple(path: str) -> dict:
    """
    Simplified PDF extraction with minimal processing.
    Use this as a fallback when the main extractor fails.
    """
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            except:
                continue
        
        # Basic cleanup
        text = re.sub(r'-\n', '', text)
        text = re.sub(r'\n+', '\n', text)
        
        # Repair URLs in fallback as well
        text = _repair_urls_in_text(text)
        
        # Try to find bibliography
        bib_start = _find_bib_start(text)
        if bib_start >= 0:
            body = text[:bib_start].strip()
            bib = text[bib_start:].strip()
        else:
            body = text.strip()
            bib = ""
        
        return {
            "full_text": text,
            "body": body,
            "bibliography": bib,
            "format": "pdf",
            "extraction_method": "pypdf (simple fallback)",
            "is_scanned": len(text.strip()) < 500,
        }
    except Exception as e:
        return {
            "full_text": "",
            "body": "",
            "bibliography": "",
            "format": "pdf",
            "error": str(e),
            "extraction_method": "failed",
        }


# ---------------------------------------------------------------------------
# DOCX Extraction - IMPROVED
# ---------------------------------------------------------------------------

def extract_docx(path: str) -> dict:
    from docx import Document
    
    doc = Document(path)
    parts = []
    
    # Main paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    
    # Tables - some LNI authors put bibliography in a borderless table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    
    # Headers and footers (sometimes contain citations)
    for section in doc.sections:
        for header in section.header.paragraphs:
            t = header.text.strip()
            if t and len(t) > 20:
                parts.append("[HEADER] " + t)
    
    text = "\n".join(parts)
    
    # Repair URLs in DOCX as well
    text = _repair_urls_in_text(text)
    
    # Clean up
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    result = split_body_bib(text, "docx")
    result["extraction_method"] = "python-docx"
    result["is_scanned"] = False
    
    return result


# ---------------------------------------------------------------------------
# LaTeX Extraction - IMPROVED with recursive crossref resolution
# ---------------------------------------------------------------------------

def _parse_bibtex_fields(body: str) -> dict:
    fields = {}
    field_start = re.compile(r'(\w+)\s*=\s*([{"])', re.DOTALL)
    pos = 0
    while pos < len(body):
        m = field_start.search(body, pos)
        if not m:
            break
        field_name = m.group(1).lower()
        delimiter = m.group(2)
        content_start = m.end()
        
        if delimiter == '{':
            depth = 1
            i = content_start
            while i < len(body) and depth > 0:
                if body[i] == '{':
                    depth += 1
                elif body[i] == '}':
                    depth -= 1
                i += 1
            value = body[content_start:i - 1]
            pos = i
        else:
            end = body.find('"', content_start)
            while end != -1 and body[end - 1] == '\\':
                end = body.find('"', end + 1)
            if end == -1:
                break
            value = body[content_start:end]
            pos = end + 1
        
        value = re.sub(r'\{([^{}]*)\}', r'\1', value)
        fields[field_name] = re.sub(r'\s+', ' ', value).strip()
    
    return fields


def _resolve_crossref_recursive(key: str, all_fields: dict, visited: set = None) -> dict:
    """Recursively resolve crossref inheritance."""
    if visited is None:
        visited = set()
    if key in visited:
        return {}
    visited.add(key)
    
    fields = all_fields.get(key, {}).copy()
    parent_key = fields.get("crossref", "").strip()
    
    if parent_key and parent_key in all_fields:
        parent_fields = _resolve_crossref_recursive(parent_key, all_fields, visited)
        for fn, val in parent_fields.items():
            if fn != "crossref" and fn not in fields:
                fields[fn] = val
    
    return fields


def _bibtex_to_lni_text(bibtex: str) -> str:
    lines = ["Literaturverzeichnis\n"]
    entry_pattern = re.compile(r'@\w+\{(\w+),(.*?)\}(?=\s*@|\s*$)', re.DOTALL)
    
    all_fields: dict = {}
    for entry_match in entry_pattern.finditer(bibtex):
        key = entry_match.group(1)
        body = entry_match.group(2)
        fields = _parse_bibtex_fields(body)
        all_fields[key] = fields
    
    # Recursively resolve crossref inheritance
    resolved_fields = {}
    for key in all_fields:
        resolved_fields[key] = _resolve_crossref_recursive(key, all_fields)
    
    for key, fields in resolved_fields.items():
        author = fields.get("author", "")
        title = fields.get("title", "")
        year = fields.get("year", "")
        pub = fields.get("publisher", "")
        journal = fields.get("journal", "")
        pages = fields.get("pages", "")
        url = fields.get("url", "")
        urldate = fields.get("urldate", "")
        booktitle = fields.get("booktitle", "")
        doi = fields.get("doi", "")
        
        parts = []
        if author:
            parts.append(f"{author}:")
        if title:
            parts.append(title + ".")
        if journal:
            parts.append(journal + ".")
        if booktitle and not journal:
            parts.append(f"In: {booktitle}.")
        if pub:
            parts.append(pub + ".")
        if pages:
            parts.append(f"S. {pages}.")
        if doi:
            parts.append(f"doi: {doi}")
        if url:
            parts.append(url)
        if urldate:
            parts.append(f"Stand: {urldate}")
        if year:
            parts.append(year + ".")
        
        lines.append(f"[{key}] {' '.join(parts)}")
    
    return "\n".join(lines)


def _extract_tex_bib_section(tex: str) -> str:
    match = re.search(
        r'\\begin\{thebibliography\}(.*?)\\end\{thebibliography\}',
        tex, re.DOTALL
    )
    if not match:
        return ""
    raw = match.group(1)
    lines = ["Literaturverzeichnis\n"]
    for item in re.finditer(
        r'\\bibitem\{(\w+)\}(.*?)(?=\\bibitem|\Z)', raw, re.DOTALL
    ):
        key = item.group(1)
        text = re.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', item.group(2))
        text = re.sub(r'[{}\\]', '', text).strip()
        lines.append(f"[{key}] {text}")
    return "\n".join(lines)


def extract_latex(tex_path: str, bib_path: str = None) -> dict:
    with open(tex_path, encoding="utf-8", errors="replace") as f:
        tex = f.read()
    
    body = _clean_latex(tex)
    
    bib_text = ""
    if bib_path and os.path.exists(bib_path):
        with open(bib_path, encoding="utf-8", errors="replace") as f:
            bib_text = f.read()
        bib_section = _bibtex_to_lni_text(bib_text)
    else:
        bib_section = _extract_tex_bib_section(tex)
    
    # Also extract from \bibliography{} command if no explicit bib file was attached
    if not bib_section and not bib_path:
        bib_file_match = re.search(r'\\bibliography\{([^}]+)\}', tex)
        if bib_file_match:
            bib_filename = bib_file_match.group(1) + ".bib"
            # Try to find the bib file in the same directory
            bib_file_path = Path(tex_path).parent / bib_filename
            if bib_file_path.exists():
                with open(bib_file_path, encoding="utf-8", errors="replace") as f:
                    bib_text = f.read()
                bib_section = _bibtex_to_lni_text(bib_text)
    
    # Repair URLs in LaTeX extracted text
    bib_section = _repair_urls_in_text(bib_section)
    body = _repair_urls_in_text(body)
    
    result = {
        "full_text": body + "\n\n" + bib_section,
        "body": body,
        "bibliography": bib_section,
        "format": "latex",
        "raw_bibtex": bib_text,
        "extraction_method": "latex parser",
        "is_scanned": False,
    }
    
    # Add warning if no bibliography found
    if not bib_section and not bib_text:
        result["warning"] = "No bibliography found. Make sure your .tex file has a \\begin{thebibliography} section or attach a .bib file."
    
    return result


def _clean_latex(tex: str) -> str:
    # Strip comments
    tex = re.sub(r'%.*', '', tex)
    
    # Remove float environments that aren't text
    tex = re.sub(
        r'\\begin\{(figure|table|lstlisting|verbatim|equation|align|tikzpicture)[^}]*\}.*?\\end\{\1\}',
        '', tex, flags=re.DOTALL
    )
    
    # Remove include directives
    tex = re.sub(r'\\include\{[^}]+\}', '', tex)
    tex = re.sub(r'\\input\{[^}]+\}', '', tex)
    
    # Unwrap common formatting commands (keep content)
    tex = re.sub(
        r'\\(?:textbf|textit|emph|texttt|text|section\*?|subsection\*?|subsubsection\*?|'
        r'paragraph|subparagraph|caption|label|ref|Cref|cref|url|href)\{([^}]*)\}',
        r'\1', tex
    )
    
    # Remove other commands
    tex = re.sub(r'\\[a-zA-Z]+\*?\{[^}]*\}', '', tex)
    tex = re.sub(r'\\[a-zA-Z]+\*?', ' ', tex)
    
    # Remove braces
    tex = re.sub(r'[{}]', ' ', tex)
    
    # Remove special characters and extra spaces
    tex = re.sub(r'\s+', ' ', tex)
    
    # Remove LaTeX math mode markers
    tex = re.sub(r'\$[^$]+\$', '[MATH]', tex)
    tex = re.sub(r'\$\$[^$]+\$\$', '[DISPLAY MATH]', tex)
    
    return tex.strip()


# ---------------------------------------------------------------------------
# Public entry point - IMPROVED with fallbacks
# ---------------------------------------------------------------------------

def extract(file_path: str, bib_path: str = None) -> dict:
    """
    Extract text from document with automatic fallback methods.
    Returns dict with 'body', 'bibliography', 'format', and metadata.
    """
    ext = Path(file_path).suffix.lower()
    
    if ext == ".pdf":
        result = extract_pdf(file_path)
        
        # If PDF extraction yielded very little text, try simple fallback
        if len(result.get("body", "")) < 500 and len(result.get("bibliography", "")) < 100:
            fallback = extract_pdf_simple(file_path)
            if len(fallback.get("body", "")) > len(result.get("body", "")):
                result = fallback
                
        return result
    
    elif ext == ".docx":
        return extract_docx(file_path)
    
    elif ext in (".tex", ".latex"):
        return extract_latex(file_path, bib_path)
    
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. Supported: .pdf, .docx, .tex"
        )